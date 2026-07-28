from pathlib import Path

from core.settings import get_settings
from processors.base import ProcessadorArquivo


class ProcessadorDOCX(ProcessadorArquivo):
    extensoes_suportadas = [".docx"]

    @classmethod
    def processar(cls, caminho: str, base_dir: Path | None = None) -> str:
        import docx

        path = cls._validar_caminho(caminho, base_dir)
        doc = docx.Document(str(path))
        texto = ""

        for paragrafo in doc.paragraphs:
            if paragrafo.text.strip():
                texto += paragrafo.text + "\n"

        for tabela in doc.tables:
            texto += "\n--- Tabela ---\n"
            for linha in tabela.rows:
                celulas = [celula.text.strip() for celula in linha.cells]
                texto += " | ".join(celulas) + "\n"
            texto += "--- Fim da Tabela ---\n"

        texto = texto.strip()
        limite_texto = get_settings().doc_text_limit
        if len(texto) > limite_texto:
            texto = texto[:limite_texto] + "\n... [Documento truncado] ..."
        return texto
