import logging
import os
from datetime import datetime

logger = logging.getLogger(__name__)


def _find_dejavu_font() -> str | None:
    """Find DejaVuSans.ttf on common system paths for Unicode PDF support."""
    candidates = [
        r"C:\Windows\Fonts\DejaVuSans.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/dejavu/DejaVuSans.ttf",
        "/System/Library/Fonts/DejaVuSans.ttf",
    ]
    for path in candidates:
        if os.path.exists(path):
            return path
    return None


class GeradorRelatorio:
    @staticmethod
    def gerar_markdown(titulo: str, conteudo: str, metadados: dict = None) -> str:
        data = datetime.now().strftime("%d/%m/%Y %H:%M")
        meta_linhas = ""
        if metadados:
            for chave, valor in metadados.items():
                meta_linhas += f"- **{chave}**: {valor}\n"

        return (
            f"# {titulo}\n\n"
            f"*Gerado em: {data}*\n\n"
            f"{meta_linhas}\n"
            f"---\n\n"
            f"{conteudo}\n"
        )

    @staticmethod
    def exportar_pdf(titulo: str, conteudo: str, caminho_saida: str, metadados: dict = None):
        from fpdf import FPDF

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        dejavu_path = _find_dejavu_font()
        if dejavu_path:
            pdf.add_font("DejaVu", "", dejavu_path, uni=True)
            pdf.add_font("DejaVu", "B", dejavu_path.replace("Sans.ttf", "Sans-Bold.ttf"), uni=True)
            pdf.add_font("DejaVu", "I", dejavu_path.replace("Sans.ttf", "Sans-Oblique.ttf"), uni=True)
            use_unicode = True
        else:
            use_unicode = False

        if use_unicode:
            pdf.set_font("DejaVu", "B", 16)
        else:
            pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 10, titulo, new_x="LMARGIN", new_y="NEXT", align="C")
        pdf.ln(5)

        if use_unicode:
            pdf.set_font("DejaVu", "", 10)
        else:
            pdf.set_font("Helvetica", "", 10)
        data = datetime.now().strftime("%d/%m/%Y %H:%M")
        pdf.cell(0, 8, f"Gerado em: {data}", new_x="LMARGIN", new_y="NEXT", align="C")
        pdf.ln(3)

        if metadados:
            if use_unicode:
                pdf.set_font("DejaVu", "I", 9)
            else:
                pdf.set_font("Helvetica", "I", 9)
            for chave, valor in metadados.items():
                pdf.cell(0, 6, f"{chave}: {valor}", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(3)

        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(5)

        if use_unicode:
            pdf.set_font("DejaVu", "", 11)
        else:
            pdf.set_font("Helvetica", "", 11)
        for linha in conteudo.split("\n"):
            if use_unicode:
                pdf.multi_cell(0, 6, linha)
            else:
                linha_limpa = linha.encode("latin-1", errors="replace").decode("latin-1")
                pdf.multi_cell(0, 6, linha_limpa)

        os.makedirs(os.path.dirname(caminho_saida) if os.path.dirname(caminho_saida) else ".", exist_ok=True)
        pdf.output(caminho_saida)
        return caminho_saida

    @staticmethod
    def exportar_docx(titulo: str, conteudo: str, caminho_saida: str, metadados: dict = None):
        import docx
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        documento = docx.Document()

        estilo = documento.styles["Normal"]
        estilo.font.name = "Calibri"
        estilo.font.size = Pt(11)

        titulo_doc = documento.add_heading(titulo, level=0)
        titulo_doc.alignment = WD_ALIGN_PARAGRAPH.CENTER

        data = datetime.now().strftime("%d/%m/%Y %H:%M")
        subtitulo = documento.add_paragraph()
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitulo.add_run(f"Gerado em: {data}")
        run.font.size = Pt(9)
        run.font.italic = True

        if metadados:
            for chave, valor in metadados.items():
                p = documento.add_paragraph()
                run_chave = p.add_run(f"{chave}: ")
                run_chave.bold = True
                p.add_run(str(valor))
            documento.add_paragraph()

        documento.add_paragraph("_" * 60)

        for linha in conteudo.split("\n"):
            linha = linha.strip()
            if not linha:
                documento.add_paragraph()
                continue

            if linha.startswith("# "):
                documento.add_heading(linha[2:], level=1)
            elif linha.startswith("## "):
                documento.add_heading(linha[3:], level=2)
            elif linha.startswith("### "):
                documento.add_heading(linha[4:], level=3)
            elif linha.startswith("- "):
                documento.add_paragraph(linha[2:], style="List Bullet")
            else:
                documento.add_paragraph(linha)

        os.makedirs(os.path.dirname(caminho_saida) if os.path.dirname(caminho_saida) else ".", exist_ok=True)
        documento.save(caminho_saida)
        return caminho_saida
